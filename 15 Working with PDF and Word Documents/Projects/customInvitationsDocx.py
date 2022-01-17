# This program reads guests names from a .txt file and 
# creates invitations in a single docx as a invitation per page.

import docx
doc = docx.Document()
file = open('customInvitationsGuests.txt')
fileContent = file.read()
fileContent = fileContent.split('\n')

for guest in fileContent:
    
    p1 = doc.add_paragraph('', 'Quote')
    p1.style.font.name = 'Monotype Corsiva'
    p1.add_run('It would be a pleasure to have the comapny of').italic = True
    p1.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    
    p1 = doc.add_paragraph('', 'Normal')
    p1.add_run(guest).bold = True
    p1.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    
    p1 = doc.add_paragraph('', 'Quote')
    p1.style.font.name = 'Monotype Corsiva'
    p1.add_run('at 11010 Memory Lane on the Evening of').italic = True
    p1.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    
    p1 = doc.add_paragraph('', 'Normal')
    p1.add_run('April 1st')
    p1.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

    p1 = doc.add_paragraph('', 'Quote')
    p1.style.font.name = 'Monotype Corsiva'
    p1.add_run("at 7 o\'clock").italic = True
    p1.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    p1.runs[0].add_break(docx.enum.text.WD_BREAK.PAGE) 
    doc.save('customInvitationsDocx.docx')
file.close()
