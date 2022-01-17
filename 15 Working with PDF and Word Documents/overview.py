import re
import PyPDF2
from PyPDF2 import pdf
from docx.api import Document

pdfFileObj = open('meetingminutes.pdf', 'rb')   # open pdf file in binary mode
pdfReader = PyPDF2.PdfFileReader(pdfFileObj)    # reads pdf file
pdfReader.numPages                              # 19    # number of pages
pageObj = pdfReader.getPage(0)
pageObj.extractText()                           # text will be displayed
pageObj.close()

# Decrypting PDFs
pdfReader = PyPDF2.PdfFileReader(open('encrypted.pdf','rb'))
pdfReader.isEncrypted                           # retrun boolean value
pdfReader.getPage(0)                            # error since file is encrypted
pdfReader = PyPDF2.PdfFileReader(open('encrypted.pdf', 'rb'))   # new pdfFileReader object because of some bug in PyPDF2
pdfReader.decrypt('rosebud')                    # 1 # passing password of file
pageObj = pdfReader.getPage(0)
    # Note: the decrypt() only decrypts the PdfFileReader object, not the the actual PDF file.
    # After your program terminates, the file on your hard drive remains encrypted.

# Creating PDFs
# Copying Pages
pdf1File = open('meetingminutes.pdf', 'rb')         # read binary
pdf2File = open('meetingminutes2.pdf', 'rb')
pdf1Reader = PyPDF2.PdfFileReader(pdf1File)
pdf2Reader = PyPDF2.PdfFileReader(pdf2File)
pdfWriter = PyPDF2.PdfFileWriter()

for pageNum in range(pdf1Reader.numPages):
    pageObj = pdf1Reader.getPage(pageNum)
    pdfWriter.addPage(pageObj)

for pageNum in range(pdf2Reader.numPages):
    pageObj = pdf2Reader.getPage(pageNum)
    pdfWriter.addPage(pageObj)

pdfOutputFile = open('combinedminutes.pdf', 'wb')   # write binary
pdfWriter.write(pdfOutputFile)
pdfOutputFile.close()
pdf1File.close()
pdf2File.close()

# Rotating Pages
minutesFile = open('meetingminutes.pdf', 'rb')
pdfReader = PyPDF2.PdfFileReader(minutesFile)
page = pdfReader.getPage(0)
page.rotateClockwise(90)
pdfWriter = PyPDF2.PdfFileWriter()
pdfWriter.addPage(page)
resultPdfFile = open('rotatedPage.pdf', 'wb')
pdfWriter.write(resultPdfFile)
resultPdfFile.close()
minutesFile.close()

# Overlaying Pages
minutesFile = open('meetingminutes.pdf', 'rb')
pdfReader = PyPDF2.PdfFileReader(minutesFile)
minutesFirstPage = pdfReader.getPage(0)
pdfWatermarkReader = PyPDF2.PdfFileReader(open('watermark.pdf', 'rb'))
minutesFirstPage.mergePage(pdfWatermarkReader.getPage(0))
pdfWriter = PyPDF2.PdfFileWriter()
pdfWriter.addPage(minutesFirstPage)

for pageNum in range(1, pdfReader.numPages):
    pageObj = pdfReader.getPage(pageNum)
    pdfWriter.addPage(pageObj)

resultPdfFile = open('watermarkedCover.pdf', 'wb')
pdfWriter.write(resultPdfFile)
minutesFile.close()
resultPdfFile.close()

# Encrypting PDFs
pdfFile = open('meetingminutes.pdf', 'rb')
pdfReader = PyPDF2.PdfFileReader(pdfFile)
pdfWriter = PyPDF2.PdfFileWriter()
for pageNum in range(pdfReader.numPages):
    pdfWriter.addPage(pdfReader.getPage(pageNum))
pdfWriter.encrypt('swordfish')
resultPdf = open('encryptedminuted.pdf', 'wb')
pdfWriter.write(resultPdf)
resultPdf.close()

'-------------------------------------------------------------------------'
# Word Documents
import docx     # docx module
doc = docx.Document('demo.docx')
len(doc.paragraphs)                 # 7
doc.paragraphs[0].text              # 'Document Title'
doc.paragraphs[1].text              # 'A plain paragraph with some bold and some italic'
len(doc.paragraphs[1].runs)         # 4
doc.paragraphs[1].runs[0].text      # 'A plain paragraph with some'
doc.paragraphs[1].runs[1].text      # 'bold'
doc.paragraphs[1].runs[2].text      # ' and some '
doc.paragraphs[1].runs[3].text      # 'italic'

# getting the full text from a .docx file
import docx
def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

# styling paragraph and run objects
'Normal'           'Heading 5'        'List Bullet'          'List Paragraph'                       
'Body Text'        'Heading 6'        'List Bullet 2'        'MacroText'                       
'Body Text 2'      'Heading 7'        'List Bullet 3'        'No Spacing'                       
'Body Text 3'      'Heading 8'        'List Continue'        'Quote'                       
'Caption'          'Heading 9'        'List Continue 2'      'Subtitle'                       
'Heading 1'        'Intense Quote'    'List Continue 3'      'TOC Heading'                       
'Heading 2'        'List'             'List Number '         'Title'                       
'Heading 3'        'List 2'           'List Number 2'                                   
'Heading 4'        'List 3'           'List Number 3'                                   

    # examples to use
#paragraphObj.style = 'Quote'
#runObj.style = 'Quote Char'     # "Char" should be added while applying to run object

# Run Attributes
'''Attribute                 Description
bold                    The text appears in bold.
italic                  The text appears in italic.
underline               The text is underlined.
strike                  The text appears with strikethrough.
double_strike           The text appears with double strikethrough.
all_caps                The text appears in capital letters.
small_caps              The text appears in capital letters, with lowercase letters two points smaller.
shadow                  The text appears with a shadow.
outline                 The text appears outlined rather than solid.
rtl                     The text is written right-to-left.
imprint                 The text appears pressed into the page.
emboss                  The text appears raised off the page in relief'''

doc.paragraphs[1].runs[0].underline = True      # usage

# Writing Word Documents
doc = docx.Document()
doc.add_paragraph('Hello, world!')
doc.save('helloworld.docx')

paraObj1 = doc.add_paragraph('This is a second paragraph.')             # add new paragraph to the docx
paraObj2 = doc.add_paragraph('This is a yet another paragraph.')        # add new paragraph to the docx
paraObj1.add_run('This text is being added to the second paragraph.')   # add text to the existing paragraph
doc.save('multiParagraphs.docx')

doc.add_paragraph('Hello, world!', 'Title') # both add_paragraph() and add_run() take second argument that is a string of Paragraph or Run object's style

# Adding headings
doc.add_heading('Header 0', 0)      # first argument is text
doc.add_heading('Header 0', 1)      # and second one is header type

# Adding line and page breaks
doc.add_paragraph('This is on first page!')
doc.paragraphs[0].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)   # docx.enum.text.WD_BREAK.PAGE  -  adds a page break
doc.add_paragraph('This is on the second page!')                    # add_break()    -   adds line break
doc.save('twoPage.docx')

# Adding pictures
doc.add_picture('zophie.png', width=docx.shared.Inches(1), height=docx.shared.Cm(4))
        # width of 1 inch and height of 4 cm
 